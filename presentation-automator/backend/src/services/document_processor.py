"""
Document processing service for handling different document formats.

This module provides utilities for extracting text from various document formats,
including DOCX, PDF, and plain text files.
"""

import asyncio
import io
import os
from pathlib import Path
from typing import Any, Dict, Optional, Tuple, Union

import docx
from PyPDF2 import PdfReader
from loguru import logger

from ..utils.context_manager import GenerationContext


class DocumentProcessor:
    """Service for processing different document formats."""
    
    def __init__(self):
        """Initialize the document processor service."""
        logger.info("Initializing document processor service")
    
    async def process_document(
        self, 
        file_path: Optional[str] = None, 
        file_content: Optional[bytes] = None,
        file_extension: Optional[str] = None,
        document_text: Optional[str] = None,
        context: Optional[GenerationContext] = None
    ) -> Tuple[str, Dict[str, Any]]:
        """
        Process a document and extract its text content.
        
        This method can process a document from:
        - A file path
        - Raw file content (bytes)
        - Direct text input
        
        Args:
            file_path: Path to the document file (optional)
            file_content: Raw content of the file as bytes (optional)
            file_extension: File extension to determine format (optional)
            document_text: Direct text input (optional)
            context: Context for tracking processing (optional)
            
        Returns:
            A tuple containing:
            - The extracted text content
            - Statistics about the document (word count, etc.)
            
        Raises:
            ValueError: If no input is provided or format is unsupported
        """
        if context:
            context.set_stage_status("document_processing", "in_progress")
        
        # Determine processing method based on input
        if document_text is not None:
            # Direct text input
            text = document_text
            stats = await self._analyze_text(text)
            
            if context:
                context.set_document_statistics(stats)
                context.set_stage_status("document_processing", "completed")
            
            return text, stats
        
        if file_path is None and file_content is None:
            raise ValueError("Either file_path, file_content, or document_text must be provided")
        
        # Determine file extension if not provided
        if file_extension is None and file_path is not None:
            file_extension = os.path.splitext(file_path)[1].lower()
        
        # Process based on file type
        if file_extension in ['.docx', '.doc']:
            text, stats = await self.extract_from_docx(file_path, file_content)
        elif file_extension in ['.pdf']:
            text, stats = await self.extract_from_pdf(file_path, file_content)
        elif file_extension in ['.txt', '.md', '.markdown', None]:
            # Handle text files or direct text content
            if file_path:
                with open(file_path, 'r', encoding='utf-8') as f:
                    text = f.read()
            else:
                text = file_content.decode('utf-8')
            
            stats = await self._analyze_text(text)
        else:
            raise ValueError(f"Unsupported file format: {file_extension}")
        
        if context:
            context.set_document_statistics(stats)
            context.set_stage_status("document_processing", "completed")
        
        return text, stats
    
    async def extract_from_docx(
        self, 
        file_path: Optional[str] = None, 
        file_content: Optional[bytes] = None
    ) -> Tuple[str, Dict[str, Any]]:
        """
        Extract text from a DOCX document.
        
        Args:
            file_path: Path to the DOCX file (optional)
            file_content: Raw content of the file as bytes (optional)
            
        Returns:
            A tuple containing:
            - The extracted text content
            - Statistics about the document
            
        Raises:
            ValueError: If neither file_path nor file_content is provided
        """
        if file_path is None and file_content is None:
            raise ValueError("Either file_path or file_content must be provided")
        
        # Create a docx Document object
        if file_path:
            doc = docx.Document(file_path)
        else:
            doc = docx.Document(io.BytesIO(file_content))
        
        # Extract text from paragraphs
        paragraphs = []
        for para in doc.paragraphs:
            if para.text.strip():
                paragraphs.append(para.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text)
        
        # Combine all text
        text = "\n\n".join(paragraphs)
        
        # Analyze the text
        stats = await self._analyze_text(text)
        
        # Add DOCX-specific statistics
        stats["document_type"] = "docx"
        stats["paragraph_count"] = len(doc.paragraphs)
        stats["table_count"] = len(doc.tables)
        
        return text, stats
    
    async def extract_from_pdf(
        self, 
        file_path: Optional[str] = None, 
        file_content: Optional[bytes] = None
    ) -> Tuple[str, Dict[str, Any]]:
        """
        Extract text from a PDF document.
        
        Args:
            file_path: Path to the PDF file (optional)
            file_content: Raw content of the file as bytes (optional)
            
        Returns:
            A tuple containing:
            - The extracted text content
            - Statistics about the document
            
        Raises:
            ValueError: If neither file_path nor file_content is provided
        """
        if file_path is None and file_content is None:
            raise ValueError("Either file_path or file_content must be provided")
        
        # Create a PDF reader object
        if file_path:
            pdf = PdfReader(file_path)
        else:
            pdf = PdfReader(io.BytesIO(file_content))
        
        # Extract text from pages
        text_parts = []
        for page in pdf.pages:
            text = page.extract_text()
            if text.strip():
                text_parts.append(text)
        
        # Combine all text
        text = "\n\n".join(text_parts)
        
        # Analyze the text
        stats = await self._analyze_text(text)
        
        # Add PDF-specific statistics
        stats["document_type"] = "pdf"
        stats["page_count"] = len(pdf.pages)
        
        return text, stats
    
    async def extract_from_text(
        self, 
        file_path: Optional[str] = None, 
        file_content: Optional[bytes] = None,
        text: Optional[str] = None
    ) -> Tuple[str, Dict[str, Any]]:
        """
        Extract and analyze text from a plain text document.
        
        Args:
            file_path: Path to the text file (optional)
            file_content: Raw content of the file as bytes (optional)
            text: Direct text input (optional)
            
        Returns:
            A tuple containing:
            - The extracted/processed text content
            - Statistics about the text
            
        Raises:
            ValueError: If no input source is provided
        """
        if text is not None:
            return text, await self._analyze_text(text)
        
        if file_path is None and file_content is None:
            raise ValueError("Either file_path, file_content, or text must be provided")
        
        # Process text file
        if file_path:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()
        else:
            text = file_content.decode('utf-8')
        
        # Analyze the text
        stats = await self._analyze_text(text)
        stats["document_type"] = "text"
        
        return text, stats
    
    async def _analyze_text(self, text: str) -> Dict[str, Any]:
        """
        Analyze text and generate statistics.
        
        Args:
            text: The text to analyze
            
        Returns:
            A dictionary containing statistics about the text
        """
        # Calculate basic statistics
        character_count = len(text)
        word_count = len(text.split())
        line_count = len(text.splitlines())
        paragraph_count = len([p for p in text.split("\n\n") if p.strip()])
        
        # Count sentences (simple approximation)
        sentence_count = text.count('.') + text.count('!') + text.count('?')
        
        # Calculate averages
        avg_words_per_paragraph = word_count / max(1, paragraph_count)
        avg_sentences_per_paragraph = sentence_count / max(1, paragraph_count)
        
        return {
            "character_count": character_count,
            "word_count": word_count,
            "line_count": line_count,
            "paragraph_count": paragraph_count,
            "sentence_count": sentence_count,
            "avg_words_per_paragraph": avg_words_per_paragraph,
            "avg_sentences_per_paragraph": avg_sentences_per_paragraph,
            "document_type": "text"
        }


# Create a singleton instance
document_processor = DocumentProcessor()

# Convenience function for direct module use
async def process_document(*args, **kwargs) -> Tuple[str, Dict[str, Any]]:
    """
    Process a document using the DocumentProcessor service.
    
    This is a convenience function that forwards all arguments to
    DocumentProcessor.process_document.
    
    Returns:
        The result from DocumentProcessor.process_document
    """
    return await document_processor.process_document(*args, **kwargs) 